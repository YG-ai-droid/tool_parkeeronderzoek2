from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "documentatie/rekenlogica-parkeeronderzoek-tool.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "D9EAF7")
        if widths:
            header_cells[i].width = Cm(widths[i])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])

    doc.add_paragraph()
    return table


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F3FF")
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(76, 29, 149)
    r.font.size = Pt(10)
    p.add_run("\n" + body).font.size = Pt(9)
    doc.add_paragraph()


def add_formula(doc, label, formula, example):
    p = doc.add_paragraph()
    p.style = "List Bullet"
    p.add_run(label + ": ").bold = True
    p.add_run(formula)
    if example:
        p.add_run(" Voorbeeld: " + example)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Heading 1", 16, "1F2937"),
        ("Heading 2", 13, "2563EB"),
        ("Heading 3", 11, "374151"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def build_doc():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Rekenlogica parkeeronderzoek-tool")
    r.font.name = "Arial"
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(17, 24, 39)

    subtitle = doc.add_paragraph()
    subtitle.add_run(
        "Controlehandleiding met formules, interpretatieregels en een volledig voorbeeld"
    ).italic = True

    add_note(
        doc,
        "Doel van dit document",
        "Dit document beschrijft hoe de tool rekent vanaf ingevoerde nummerplaten tot de "
        "analysekaarten, taartdiagrammen, rotatieanalyse, herhalingsanalyse en parkeerprofielen. "
        "Het is opgezet als controlepad: eerst de brondata, daarna elke afgeleide stap.",
    )

    doc.add_heading("1. Basisgegevens waarop alle berekeningen steunen", level=1)
    doc.add_paragraph(
        "De tool rekent vanuit drie hoofdtabellen: telmomenten, zones en tellingen per zone. "
        "Clusters en profielen zijn afgeleide structuren die verwijzen naar zones en telmomenten."
    )

    add_table(
        doc,
        ["Onderdeel", "Betekenis", "Belangrijk voor de berekening"],
        [
            ["Telmoment", "Naam, datum en tijdstip van een telling.", "Bepaalt in welke kolom of periode een registratie valt."],
            ["Zone", "Tellocatie met capaciteit, parkeerregime en optionele polygonen.", "Capaciteit is de noemer voor bezettingsgraad."],
            ["Telling", "Lijst nummerplaten per zone per telmoment.", "Elke unieke nummerplaat in die lijst telt als 1 bezette plaats."],
            ["Cluster", "Groep van zones.", "Capaciteit en aantal voertuigen worden over de clusterzones opgeteld."],
            ["Parkeerprofiel", "Regelset met tijdvenster en minimale/maximale duur.", "Filtert voertuigsessies die aan het profiel voldoen."],
        ],
        widths=[3.0, 6.2, 7.0],
    )

    doc.add_heading("2. Voorbeelddata voor controle", level=1)
    doc.add_paragraph(
        "Alle voorbeelden in dit document gebruiken dezelfde kleine dataset. Zo kan elke berekening "
        "handmatig worden nagerekend."
    )

    add_table(
        doc,
        ["Telmoment", "Datum", "Tijdstip"],
        [
            ["T1", "15/01/2026", "08:00"],
            ["T2", "15/01/2026", "12:00"],
            ["T3", "15/01/2026", "18:00"],
            ["T4", "16/01/2026", "08:00"],
        ],
        widths=[3.2, 4.0, 3.2],
    )

    add_table(
        doc,
        ["Zone", "Capaciteit", "Regime"],
        [
            ["Zone A", "4", "vrij parkeren"],
            ["Zone B", "3", "blauwe zone"],
            ["Zone C", "5", "bewonerskaart"],
        ],
        widths=[4.0, 3.0, 5.5],
    )

    add_table(
        doc,
        ["Zone", "T1 08:00", "T2 12:00", "T3 18:00", "T4 08:00"],
        [
            ["A", "AAA, BBB", "AAA, CCC, DDD", "AAA, DDD, EEE, FFF", "BBB, EEE"],
            ["B", "GGG", "GGG, HHH", "HHH, III", "HHH"],
            ["C", "JJJ, KKK", "JJJ", "LLL, MMM, NNN", "JJJ, LLL"],
        ],
        widths=[1.8, 4.1, 4.1, 4.3, 3.8],
    )

    add_note(
        doc,
        "Cluster in het voorbeeld",
        "Cluster Centrum bestaat uit Zone A en Zone B. Zone C blijft buiten deze cluster.",
    )

    doc.add_heading("3. Normalisatie en privacy", level=1)
    doc.add_paragraph(
        "Voor de berekeningen worden nummerplaten genormaliseerd: hoofdletters, geen spaties of "
        "streepjes, en een maximumlengte. Binnen eenzelfde zone en telmoment bewaart de tool elke "
        "nummerplaat maar een keer. Bij versleuteling wordt dezelfde genormaliseerde nummerplaat "
        "telkens naar dezelfde code omgezet binnen hetzelfde project."
    )
    add_formula(
        doc,
        "Normalisatie",
        "AB-123 wordt AB123; ab 123 wordt eveneens AB123.",
        "Daardoor herkent de tool dit als hetzelfde voertuig.",
    )
    add_formula(
        doc,
        "Versleuteling",
        "ENC + eerste 16 hex-tekens van SHA-256(projectsalt + nummerplaat).",
        "AB123 krijgt in hetzelfde project altijd dezelfde ENC-code.",
    )

    doc.add_heading("4. Bezetting per zone en telmoment", level=1)
    add_formula(
        doc,
        "Aantal voertuigen",
        "lengte van de nummerplatenlijst voor zone en telmoment.",
        "Zone A op T2 heeft AAA, CCC, DDD = 3 voertuigen.",
    )
    add_formula(
        doc,
        "Bezettingsgraad",
        "round(aantal voertuigen / capaciteit * 100).",
        "Zone A op T2: round(3 / 4 * 100) = 75%.",
    )

    add_table(
        doc,
        ["Zone", "T2 aantal", "Capaciteit", "Bezettingsgraad"],
        [
            ["A", "3", "4", "75%"],
            ["B", "2", "3", "67%"],
            ["C", "1", "5", "20%"],
        ],
        widths=[3.0, 3.0, 3.0, 4.0],
    )

    doc.add_heading("5. Kleurcodes", level=1)
    doc.add_paragraph(
        "De kleurcode wordt bepaald op basis van de bezettingsgraad en de ingestelde grenzen. "
        "De standaardgrenzen zijn lichtgrijs tot 40%, groen tot 70% en oranje tot 85%. "
        "De tool gebruikt ondergrenzen impliciet en vergelijkt met 'kleiner dan'."
    )
    add_table(
        doc,
        ["Voorwaarde", "Kleur", "Voorbeeld"],
        [
            ["bezetting < lichtgrijsTot", "lichtgrijs", "20% bij Zone C op T2"],
            ["bezetting < groenTot", "groen", "67% bij Zone B op T2"],
            ["bezetting < oranjeTot", "oranje", "75% bij Zone A op T2"],
            ["anders", "rood", "90% of hoger bij standaardgrenzen"],
        ],
        widths=[5.2, 3.0, 6.8],
    )

    add_note(
        doc,
        "Controlepunt",
        "Als een grens exact wordt geraakt, schuift de waarde naar de volgende kleur. "
        "Bij standaardgrenzen is 70% dus oranje, niet groen.",
    )

    doc.add_heading("6. Analyse: alle zones op gekozen telmoment", level=1)
    doc.add_paragraph(
        "Deze analyse telt alle geselecteerde zones op voor een gekozen telmoment. De totale "
        "bezettingsgraad gebruikt de som van de voertuigen en de som van de capaciteit."
    )
    add_formula(
        doc,
        "Totale voertuigen",
        "som voertuigen over alle analysezones.",
        "T2: A 3 + B 2 + C 1 = 6.",
    )
    add_formula(
        doc,
        "Totale capaciteit",
        "som capaciteit over alle analysezones.",
        "A 4 + B 3 + C 5 = 12.",
    )
    add_formula(
        doc,
        "Totale bezetting",
        "round(totale voertuigen / totale capaciteit * 100).",
        "round(6 / 12 * 100) = 50%.",
    )
    add_formula(
        doc,
        "Drukste telmoment",
        "telmoment met de hoogste som voertuigen over de analysezones.",
        "T1=5, T2=6, T3=9, T4=5; dus T3 is drukst.",
    )

    doc.add_heading("7. Gemiddelde over alle telmomenten", level=1)
    doc.add_paragraph(
        "Bij 'gemiddelde over alle telmomenten' wordt eerst per zone het gemiddelde aantal "
        "voertuigen berekend. Daarna wordt dit gedeeld door de capaciteit van die zone."
    )
    add_formula(
        doc,
        "Gemiddeld aantal zone",
        "som aantallen over telmomenten / aantal telmomenten.",
        "Zone A: (2 + 3 + 4 + 2) / 4 = 2,75.",
    )
    add_formula(
        doc,
        "Gemiddelde bezetting zone",
        "round(gemiddeld aantal / capaciteit * 100).",
        "Zone A: round(2,75 / 4 * 100) = 69%.",
    )

    doc.add_heading("8. Clusters", level=1)
    doc.add_paragraph(
        "Een cluster is een verzameling zones. Voor bezetting telt de tool de capaciteit en het "
        "aantal voertuigen van de zones in de cluster op."
    )
    add_formula(
        doc,
        "Cluster capaciteit",
        "som capaciteit van clusterzones.",
        "Centrum: Zone A 4 + Zone B 3 = 7.",
    )
    add_formula(
        doc,
        "Cluster aantal",
        "som voertuigen in clusterzones voor het gekozen telmoment.",
        "Centrum op T2: Zone A 3 + Zone B 2 = 5.",
    )
    add_formula(
        doc,
        "Cluster bezetting",
        "round(cluster aantal / cluster capaciteit * 100).",
        "round(5 / 7 * 100) = 71%.",
    )
    add_note(
        doc,
        "Interpretatieverschil",
        "Voor gewone bezetting telt de tool registraties per zone op. Voor de herhalingsanalyse "
        "van dezelfde voertuigen telt de nummerplaat zelf als voertuig, zodat dezelfde plaat binnen "
        "een cluster niet kunstmatig als meerdere voertuigen wordt gezien.",
    )

    doc.add_heading("9. Rotatieanalyse", level=1)
    doc.add_paragraph(
        "De rotatieanalyse kijkt naar opeenvolgende telmomenten waarin dezelfde nummerplaat in "
        "dezelfde analyse-eenheid voorkomt. De analyse-eenheid is een zone of een cluster."
    )
    add_table(
        doc,
        ["Nummerplaat Zone A", "Aanwezig in", "Reeks(en)", "Interpretatie"],
        [
            ["AAA", "T1, T2, T3", "T1-T3", "een voertuig bleef drie opeenvolgende telrondes"],
            ["BBB", "T1, T4", "T1-T1 en T4-T4", "geen aaneengesloten reeks"],
            ["DDD", "T2, T3", "T2-T3", "een voertuig bleef twee opeenvolgende telrondes"],
            ["EEE", "T3, T4", "T3-T4", "loopt door van avond naar volgende ochtend"],
        ],
        widths=[3.8, 3.4, 3.6, 6.2],
    )
    doc.add_paragraph(
        "De blauwe balkjes in de rotatieanalyse groeperen voertuigen met dezelfde start- en "
        "eindtelling. Het cijfer in een balk is het aantal voertuigen in die reeks."
    )

    doc.add_heading("10. Registratie van dezelfde voertuigen over tellingen heen", level=1)
    doc.add_paragraph(
        "Deze analyse telt per zone hoeveel verschillende nummerplaten 1 keer, 2 keer, 3 keer, "
        "enzovoort gezien zijn over alle telmomenten. De laatste categorie is instelbaar door de "
        "analist, bijvoorbeeld '4 of meer keer' of '5 of meer keer'."
    )
    add_table(
        doc,
        ["Zone A nummerplaat", "Aantal telmomenten gezien"],
        [
            ["AAA", "3"],
            ["BBB", "2"],
            ["CCC", "1"],
            ["DDD", "2"],
            ["EEE", "2"],
            ["FFF", "1"],
        ],
        widths=[5.0, 5.0],
    )
    add_table(
        doc,
        ["Categorie bij laatste categorie vanaf 4", "Aantal unieke voertuigen"],
        [
            ["1 keer", "2 (CCC, FFF)"],
            ["2 keer", "3 (BBB, DDD, EEE)"],
            ["3 keer", "1 (AAA)"],
            ["4 of meer keer", "0"],
        ],
        widths=[6.5, 7.5],
    )

    doc.add_heading("11. Registratie van dezelfde voertuigen per teldag", level=1)
    doc.add_paragraph(
        "In de analyse 'alle telmomenten voor zone of cluster' maakt de tool per teldag een "
        "soortgelijk taartdiagram. Alleen de telmomenten op die datum worden dan meegenomen."
    )
    add_formula(
        doc,
        "Per teldag",
        "groepeer telmomenten op datum, tel per nummerplaat in die daggroep.",
        "Zone A op 15/01/2026 gebruikt T1, T2 en T3.",
    )
    add_table(
        doc,
        ["Zone A op 15/01/2026", "Aantal binnen die dag"],
        [
            ["AAA", "3"],
            ["BBB", "1"],
            ["CCC", "1"],
            ["DDD", "2"],
            ["EEE", "1"],
            ["FFF", "1"],
        ],
        widths=[5.0, 5.0],
    )
    add_table(
        doc,
        ["Categorie", "Aantal unieke voertuigen"],
        [
            ["1 keer", "4 (BBB, CCC, EEE, FFF)"],
            ["2 keer", "1 (DDD)"],
            ["3 keer", "1 (AAA)"],
            ["4 of meer keer", "0"],
        ],
        widths=[5.0, 7.0],
    )

    doc.add_heading("12. Parkeerprofielen", level=1)
    doc.add_paragraph(
        "Parkeerprofielen werken niet met losse tellingen, maar met voertuigsessies. Een sessie "
        "is een aaneengesloten reeks telmomenten waarin dezelfde nummerplaat in dezelfde zone "
        "voorkomt."
    )
    add_table(
        doc,
        ["Profielregel", "Betekenis"],
        [
            ["Tijdvenster", "Welke telmomenten mogen als relevante momenten van de dag tellen."],
            ["Minimale duur", "Minimum aantal aaneengesloten telrondes in de sessie."],
            ["Maximale duur", "Optioneel maximum aantal aaneengesloten telrondes."],
            ["Resultaat per zone", "Aantal unieke nummerplaten waarvan minstens een sessie aan het profiel voldoet."],
        ],
        widths=[4.0, 10.5],
    )
    add_note(
        doc,
        "Voorbeeld bewonersprofiel",
        "Profiel 'bewoner': tijdvenster 18:00 tot 08:00, minimale duur 2, geen maximum. "
        "De relevante groep is T3 op 15/01 en T4 op 16/01. Zone A heeft EEE in T3 en T4, "
        "dus Zone A krijgt 1 voertuig voor dit profiel. Zone C heeft LLL in T3 en T4, dus Zone C krijgt ook 1.",
    )

    doc.add_heading("13. Excel-import en effect op berekeningen", level=1)
    doc.add_paragraph(
        "Het Excel-invulblad bevat per project, telmoment, tellocatie, capaciteit, parkeerregime, "
        "parkeervak en nummerplaat een rij. Bij import worden alleen ingevulde nummerplaten "
        "overgenomen. De tool koppelt op telmomentgegevens en tellocatiegegevens."
    )
    add_table(
        doc,
        ["Stap", "Controle"],
        [
            ["Lege cellen", "Worden genegeerd."],
            ["Dubbele plaat binnen zelfde zone en telmoment", "Wordt maar een keer toegevoegd."],
            ["Bestaande registratie", "Wordt niet opnieuw toegevoegd."],
            ["Privacyknop", "Vervangt leesbare platen door stabiele versleutelde codes."],
        ],
        widths=[5.0, 9.5],
    )

    doc.add_heading("14. Checklist om rekenlogica handmatig te controleren", level=1)
    checks = [
        "Kies een telmoment en tel per zone het aantal nummerplaten.",
        "Deel per zone door de capaciteit en rond naar een geheel percentage.",
        "Tel voor het totaaloverzicht eerst voertuigen en capaciteit op, en deel pas daarna.",
        "Controleer bij gemiddelde analyses dat eerst per zone het gemiddelde aantal wordt berekend.",
        "Controleer bij clusters of de juiste zones in de cluster zitten.",
        "Controleer rotatie door per nummerplaat aaneengesloten reeksen van telmomenten te markeren.",
        "Controleer herhaling door per nummerplaat het aantal telmomenten te tellen.",
        "Controleer per teldag dat alleen telmomenten met dezelfde datum meetellen.",
        "Controleer parkeerprofielen door eerst sessies te maken en daarna de profielregels toe te passen.",
    ]
    for check in checks:
        p = doc.add_paragraph(style="List Number")
        p.add_run(check)

    doc.add_heading("15. Waar deze logica in de code zit", level=1)
    add_table(
        doc,
        ["Onderdeel", "Bestand"],
        [
            ["Basisdata, zones, clusters, import/export, kleuren", "src/App.jsx"],
            ["Analyses, gemiddelden, herhaling, profielen", "src/components/AnalistDashboard.jsx"],
            ["Rotatieanalyse", "src/components/RotatieAnalyse.jsx"],
            ["Kaartkleuren en kaartpopups", "src/components/ParkeerKaart.jsx"],
        ],
        widths=[6.0, 8.0],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Bijlage: korte formulekaart", level=1)
    add_table(
        doc,
        ["Berekening", "Formule"],
        [
            ["Zonebezetting", "round(aantal voertuigen / capaciteit * 100)"],
            ["Totaalbezetting", "round(som voertuigen / som capaciteit * 100)"],
            ["Gemiddeld aantal zone", "som aantallen over telmomenten / aantal telmomenten"],
            ["Cluster aantal", "som voertuigen van de clusterzones"],
            ["Cluster capaciteit", "som capaciteit van de clusterzones"],
            ["Rotatie", "groep per nummerplaat opeenvolgende telmoment-indexen"],
            ["Herhaling", "tel per nummerplaat het aantal telmomenten waarin ze voorkomt"],
            ["Profiel", "filter sessies op tijdvenster en min/max duur"],
        ],
        widths=[5.0, 9.5],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
